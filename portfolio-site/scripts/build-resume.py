from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import HRFlowable, KeepTogether, Paragraph, SimpleDocTemplate, Spacer


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "public" / "documents"
OUTPUT_PATH = OUTPUT_DIR / "Leo-Mabuku-Resume.docx"
PDF_PATH = OUTPUT_DIR / "Leo-Mabuku-Resume.pdf"

# compact_reference_guide preset, with a named one-page ATS override:
# Calibri, compact spacing, blue hierarchy, real Word headings/lists, no tables.
NAVY = RGBColor(31, 77, 120)
BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(32, 38, 43)
MUTED = RGBColor(84, 92, 99)
RULE = "B8C7D5"
FONT = "Calibri"


SUMMARY = (
    "Fourth-year BSc Computer Science student and software developer building Android, web, systems and programming-language projects. "
    "Combines C/C++, Python and Kotlin development with entrepreneurship, live network operations, testing and user-focused problem solving. "
    "Open to graduate, internship and junior software engineering opportunities."
)


def set_cell_safe_font(run, size=10.0, bold=None, color=DARK, italic=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_hyperlink(paragraph, text, url, *, size=9.2, color=NAVY, bold=False):
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    properties.append(fonts)
    colour = OxmlElement("w:color")
    colour.set(qn("w:val"), f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
    properties.append(colour)
    size_el = OxmlElement("w:sz")
    size_el.set(qn("w:val"), str(int(size * 2)))
    properties.append(size_el)
    if bold:
        properties.append(OxmlElement("w:b"))
    run.append(properties)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_bottom_border(paragraph, color=RULE, size="8", space="5"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def configure_document(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.56)
    section.bottom_margin = Inches(0.52)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)

    normal = document.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(9.65)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(2.4)
    normal.paragraph_format.line_spacing = 1.04

    for style_name, size, color, before, after in [
        ("Heading 1", 10.4, BLUE, 6.0, 3.2),
        ("Heading 2", 10.0, NAVY, 3.0, 1.2),
        ("Heading 3", 9.7, NAVY, 2.0, 1.0),
    ]:
        style = document.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    bullet = document.styles["List Bullet"]
    bullet.font.name = FONT
    bullet._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    bullet._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    bullet.font.size = Pt(9.45)
    bullet.font.color.rgb = DARK
    bullet.paragraph_format.left_indent = Inches(0.28)
    bullet.paragraph_format.first_line_indent = Inches(-0.14)
    bullet.paragraph_format.space_before = Pt(0)
    bullet.paragraph_format.space_after = Pt(1.1)
    bullet.paragraph_format.line_spacing = 1.0

    document.core_properties.title = "Leo Mabuku Resume"
    document.core_properties.subject = "Software Developer, Computer Science Student and Entrepreneur"
    document.core_properties.author = "Leo Mabuku"
    document.core_properties.keywords = "software developer, computer science, Kotlin, Python, C, C++, Android, Zambia"


def section_heading(document, text):
    paragraph = document.add_paragraph(style="Heading 1")
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text.upper())
    set_cell_safe_font(run, size=10.4, bold=True, color=BLUE)
    add_bottom_border(paragraph, size="5", space="3")
    return paragraph


def add_label_line(document, label, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(1.2)
    label_run = paragraph.add_run(f"{label}: ")
    set_cell_safe_font(label_run, size=9.4, bold=True, color=NAVY)
    value_run = paragraph.add_run(text)
    set_cell_safe_font(value_run, size=9.4, color=DARK)
    return paragraph


def entry_heading(document, title, organisation, right_text):
    paragraph = document.add_paragraph(style="Heading 2")
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(7.14), WD_TAB_ALIGNMENT.RIGHT)
    title_run = paragraph.add_run(title)
    set_cell_safe_font(title_run, size=10.0, bold=True, color=NAVY)
    organisation_run = paragraph.add_run(f" | {organisation}")
    set_cell_safe_font(organisation_run, size=9.7, bold=True, color=DARK)
    right_run = paragraph.add_run(f"\t{right_text}")
    set_cell_safe_font(right_run, size=9.0, bold=True, color=MUTED)
    return paragraph


def add_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.keep_together = True
    run = paragraph.add_run(text)
    set_cell_safe_font(run, size=9.45, color=DARK)
    return paragraph


def add_project(document, name, stack, url, description):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(1.2)
    paragraph.paragraph_format.space_after = Pt(0.8)
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.keep_with_next = False
    add_hyperlink(paragraph, name, url, size=9.7, color=NAVY, bold=True)
    stack_run = paragraph.add_run(f" | {stack} — ")
    set_cell_safe_font(stack_run, size=9.25, bold=True, color=MUTED)
    description_run = paragraph.add_run(description)
    set_cell_safe_font(description_run, size=9.25, color=DARK)
    return paragraph


def build_resume():
    document = Document()
    configure_document(document)

    # customer_pack-inspired left-aligned identity header, kept ATS-readable.
    name = document.add_paragraph()
    name.paragraph_format.space_after = Pt(0)
    name_run = name.add_run("LEO MABUKU")
    set_cell_safe_font(name_run, size=25.5, bold=True, color=NAVY)

    role = document.add_paragraph()
    role.paragraph_format.space_after = Pt(1.8)
    role_run = role.add_run("SOFTWARE DEVELOPER | COMPUTER SCIENCE STUDENT | ENTREPRENEUR")
    set_cell_safe_font(role_run, size=10.0, bold=True, color=BLUE)

    contact = document.add_paragraph()
    contact.paragraph_format.space_after = Pt(4.2)
    contact.add_run("Kitwe, Zambia | +260 770 464 736 | ")
    add_hyperlink(contact, "leokmabuku@gmail.com", "mailto:leokmabuku@gmail.com")
    contact.add_run(" | ")
    add_hyperlink(contact, "Portfolio", "https://leo-mabuku-portfolio.pages.dev")
    contact.add_run(" | ")
    add_hyperlink(contact, "GitHub", "https://github.com/leomabuku")
    contact.add_run(" | ")
    add_hyperlink(contact, "LinkedIn", "https://www.linkedin.com/in/leo-mabuku-7a2645330/")
    for run in contact.runs:
        set_cell_safe_font(run, size=9.15, color=MUTED)
    add_bottom_border(contact, size="8", space="4")

    section_heading(document, "Professional Summary")
    summary = document.add_paragraph()
    summary.paragraph_format.space_after = Pt(1.2)
    summary_run = summary.add_run(SUMMARY)
    set_cell_safe_font(summary_run, size=9.55, color=DARK)

    section_heading(document, "Technical Skills")
    add_label_line(document, "Languages", "C, C++, Python, Kotlin, Java, JavaScript, SQL, HTML and CSS")
    add_label_line(document, "Mobile & web", "Android Studio, Jetpack Compose, Room/SQLite, Firebase, Node.js, Express and Astro")
    add_label_line(document, "Systems & delivery", "Git/GitHub, Pytest, PLY, Tkinter, PostgreSQL, Cloudflare Pages and technical documentation")
    add_label_line(document, "Additional", "Networking fundamentals, hardware diagnostics and repair, component replacement and soldering")

    section_heading(document, "Experience")
    entry_heading(document, "Founder & Network Service Operator", "Starlink Internet Service", "Kitwe | Feb 2026–Present")
    add_bullet(document, "Operate a Starlink Mini and ARRIS mesh service for 29 subscribers, covering a boarding house and approximately 30 metres beyond the property.")
    add_bullet(document, "Manage onboarding, payments, subscription periods and connectivity support; designed SubTrack BH to centralise operational records.")

    entry_heading(document, "Founder, Sales & Operations", "Leo M. Suppliers", "Livingstone | Feb 2022–Apr 2026")
    add_bullet(document, "Developed a field-based bedding venture into a PACRA-registered business with a physical shop opened in September 2022.")
    add_bullet(document, "Managed procurement, stock, sales, marketing, customer service and deliveries; adapted the product range to furniture and learned assembly.")

    entry_heading(document, "Independent Software Development & Technical Tutoring", "Self-employed", "2024–Present")
    add_bullet(document, "Provide paid programming tutoring, debugging, project setup and feature customisation, with an emphasis on explaining code and project structure.")

    entry_heading(document, "Data Analytics Trainee", "MedTourEasy", "Remote | Sep 2025")
    add_bullet(document, "Completed a four-week training programme and live analytics project, applying structured analysis while observing professional confidentiality.")

    section_heading(document, "Selected Projects")
    add_project(
        document,
        "TongaLang",
        "Python, PLY, Tkinter, Pytest",
        "https://github.com/leomabuku/PROJECT",
        "Educational interpreted language with a custom lexer, parser, AST, interpreter, IDE tools, bilingual errors, 17 examples and 147 passing automated tests.",
    )
    add_project(
        document,
        "SubTrack BH",
        "Kotlin, Compose, Room, Firebase",
        "https://github.com/leomabuku/SubTrackBH",
        "Local-first subscriber and payment workspace designed around the active 29-subscriber internet service.",
    )
    add_project(
        document,
        "CBU-FIND",
        "Kotlin, Next.js, Firebase",
        "https://cbu-find-web.leokmabuku.workers.dev",
        "Cross-platform campus lost-and-found with authentication, image reports, search and private item-linked conversations.",
    )
    add_project(
        document,
        "SERC Mini-OS",
        "C, Raylib, Win32",
        "https://github.com/leomabuku/SERC-Mini-OS-system",
        "Operating-systems simulator covering four CPU schedulers, memory allocation, paging, IPC, deadlock safety and regression tests.",
    )

    section_heading(document, "Education & Professional Development")
    entry_heading(document, "BSc Computer Science", "The Copperbelt University", "2023–Present | Expected Sep 2026")
    add_label_line(document, "Training", "Prompt Engineering for Everyone — IBM Skills Network / Cognitive Class, Oct 2025; Data Analytics Traineeship — MedTourEasy, Sep 2025")
    add_label_line(document, "Business courses", "Marketing & Sales, Money Skills and Entrepreneurial Skills — Absa Skills")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


def build_pdf():
    regular_font = "Helvetica"
    bold_font = "Helvetica-Bold"
    font_dir = Path("C:/Windows/Fonts")
    regular_path = font_dir / "calibri.ttf"
    bold_path = font_dir / "calibrib.ttf"
    if regular_path.exists() and bold_path.exists():
        pdfmetrics.registerFont(TTFont("CalibriResume", str(regular_path)))
        pdfmetrics.registerFont(TTFont("CalibriResume-Bold", str(bold_path)))
        regular_font = "CalibriResume"
        bold_font = "CalibriResume-Bold"

    navy = colors.HexColor("#1F4D78")
    blue = colors.HexColor("#2E74B5")
    dark = colors.HexColor("#20262B")
    muted = colors.HexColor("#545C63")
    rule = colors.HexColor("#B8C7D5")

    document = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        rightMargin=0.66 * inch,
        leftMargin=0.66 * inch,
        topMargin=0.55 * inch,
        bottomMargin=0.52 * inch,
        title="Leo Mabuku Resume",
        author="Leo Mabuku",
        subject="Software Developer, Computer Science Student and Entrepreneur",
    )

    styles = {
        "name": ParagraphStyle("ResumeName", fontName=bold_font, fontSize=24, leading=25, textColor=navy, spaceAfter=0),
        "role": ParagraphStyle("ResumeRole", fontName=bold_font, fontSize=10.1, leading=11.8, textColor=blue, spaceAfter=2.0),
        "contact": ParagraphStyle("ResumeContact", fontName=regular_font, fontSize=8.7, leading=10.0, textColor=muted, spaceAfter=3.2),
        "section": ParagraphStyle("ResumeSection", fontName=bold_font, fontSize=10.0, leading=11.0, textColor=blue, spaceBefore=4.7, spaceAfter=1.0),
        "body": ParagraphStyle("ResumeBody", fontName=regular_font, fontSize=9.2, leading=10.8, textColor=dark, spaceAfter=1.4),
        "label": ParagraphStyle("ResumeLabel", fontName=regular_font, fontSize=9.0, leading=10.5, textColor=dark, spaceAfter=0.7),
        "entry": ParagraphStyle("ResumeEntry", fontName=regular_font, fontSize=9.15, leading=10.4, textColor=dark, spaceBefore=1.8, spaceAfter=0.45),
        "bullet": ParagraphStyle("ResumeBullet", fontName=regular_font, fontSize=8.9, leading=10.2, textColor=dark, leftIndent=10, firstLineIndent=-7, bulletIndent=0, spaceAfter=0.55),
        "project": ParagraphStyle("ResumeProject", fontName=regular_font, fontSize=8.9, leading=10.2, textColor=dark, spaceBefore=0.8, spaceAfter=0.3),
    }

    def p(text, style):
        return Paragraph(text, styles[style])

    def bullet(text):
        return Paragraph(text, styles["bullet"], bulletText="•")

    def section(title):
        return [p(title.upper(), "section"), HRFlowable(width="100%", thickness=0.55, color=rule, spaceBefore=0, spaceAfter=1.3)]

    story = [
        p("LEO MABUKU", "name"),
        p("SOFTWARE DEVELOPER | COMPUTER SCIENCE STUDENT | ENTREPRENEUR", "role"),
        p(
            'Kitwe, Zambia | +260 770 464 736 | '
            '<link href="mailto:leokmabuku@gmail.com" color="#1F4D78">leokmabuku@gmail.com</link> | '
            '<link href="https://leo-mabuku-portfolio.pages.dev" color="#1F4D78">Portfolio</link> | '
            '<link href="https://github.com/leomabuku" color="#1F4D78">GitHub</link> | '
            '<link href="https://www.linkedin.com/in/leo-mabuku-7a2645330/" color="#1F4D78">LinkedIn</link>',
            "contact",
        ),
        HRFlowable(width="100%", thickness=0.8, color=navy, spaceBefore=0, spaceAfter=0.8),
    ]

    story += section("Professional Summary")
    story.append(p(SUMMARY, "body"))

    story += section("Technical Skills")
    story.extend([
        p(f'<font name="{bold_font}" color="#1F4D78">Languages:</font> C, C++, Python, Kotlin, Java, JavaScript, SQL, HTML and CSS', "label"),
        p(f'<font name="{bold_font}" color="#1F4D78">Mobile &amp; web:</font> Android Studio, Jetpack Compose, Room/SQLite, Firebase, Node.js, Express and Astro', "label"),
        p(f'<font name="{bold_font}" color="#1F4D78">Systems &amp; delivery:</font> Git/GitHub, Pytest, PLY, Tkinter, PostgreSQL, Cloudflare Pages and technical documentation', "label"),
        p(f'<font name="{bold_font}" color="#1F4D78">Additional:</font> Networking fundamentals, hardware diagnostics and repair, component replacement and soldering', "label"),
    ])

    story += section("Experience")
    story.append(KeepTogether([
        p(f'<font name="{bold_font}" color="#1F4D78">Founder &amp; Network Service Operator</font> | <font name="{bold_font}">Starlink Internet Service</font> — Kitwe | Feb 2026–Present', "entry"),
        bullet("Operate a Starlink Mini and ARRIS mesh service for 29 subscribers, covering a boarding house and approximately 30 metres beyond the property."),
        bullet("Manage onboarding, payments, subscription periods and connectivity support; designed SubTrack BH to centralise operational records."),
    ]))
    story.append(KeepTogether([
        p(f'<font name="{bold_font}" color="#1F4D78">Founder, Sales &amp; Operations</font> | <font name="{bold_font}">Leo M. Suppliers</font> — Livingstone | Feb 2022–Apr 2026', "entry"),
        bullet("Developed a field-based bedding venture into a PACRA-registered business with a physical shop opened in September 2022."),
        bullet("Managed procurement, stock, sales, marketing, customer service and deliveries; adapted the product range to furniture and learned assembly."),
    ]))
    story.append(KeepTogether([
        p(f'<font name="{bold_font}" color="#1F4D78">Independent Software Development &amp; Technical Tutoring</font> | <font name="{bold_font}">Self-employed</font> — 2024–Present', "entry"),
        bullet("Provide paid programming tutoring, debugging, project setup and feature customisation, with an emphasis on explaining code and project structure."),
    ]))
    story.append(KeepTogether([
        p(f'<font name="{bold_font}" color="#1F4D78">Data Analytics Trainee</font> | <font name="{bold_font}">MedTourEasy</font> — Remote | Sep 2025', "entry"),
        bullet("Completed a four-week training programme and live analytics project, applying structured analysis while observing professional confidentiality."),
    ]))

    story += section("Selected Projects")
    story.extend([
        p(f'<link href="https://github.com/leomabuku/PROJECT" color="#1F4D78"><font name="{bold_font}">TongaLang</font></link> | <font name="{bold_font}" color="#545C63">Python, PLY, Tkinter, Pytest</font> — Educational interpreted language with a custom lexer, parser, AST, interpreter, IDE tools, bilingual errors, 17 examples and 147 passing automated tests.', "project"),
        p(f'<link href="https://github.com/leomabuku/SubTrackBH" color="#1F4D78"><font name="{bold_font}">SubTrack BH</font></link> | <font name="{bold_font}" color="#545C63">Kotlin, Compose, Room, Firebase</font> — Local-first subscriber and payment workspace designed around the active 29-subscriber internet service.', "project"),
        p(f'<link href="https://cbu-find-web.leokmabuku.workers.dev" color="#1F4D78"><font name="{bold_font}">CBU-FIND</font></link> | <font name="{bold_font}" color="#545C63">Kotlin, Next.js, Firebase</font> — Cross-platform campus lost-and-found with authentication, image reports, search and private item-linked conversations.', "project"),
        p(f'<link href="https://github.com/leomabuku/SERC-Mini-OS-system" color="#1F4D78"><font name="{bold_font}">SERC Mini-OS</font></link> | <font name="{bold_font}" color="#545C63">C, Raylib, Win32</font> — Operating-systems simulator covering four CPU schedulers, memory allocation, paging, IPC, deadlock safety and regression tests.', "project"),
    ])

    story += section("Education & Professional Development")
    story.extend([
        p(f'<font name="{bold_font}" color="#1F4D78">BSc Computer Science</font> | <font name="{bold_font}">The Copperbelt University</font> — 2023–Present | Expected Sep 2026', "entry"),
        p(f'<font name="{bold_font}" color="#1F4D78">Training:</font> Prompt Engineering for Everyone — IBM Skills Network / Cognitive Class, Oct 2025; Data Analytics Traineeship — MedTourEasy, Sep 2025', "label"),
        p(f'<font name="{bold_font}" color="#1F4D78">Business courses:</font> Marketing &amp; Sales, Money Skills and Entrepreneurial Skills — Absa Skills', "label"),
    ])

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document.build(story)
    print(PDF_PATH)


if __name__ == "__main__":
    build_resume()
    build_pdf()
